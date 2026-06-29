"""Dummy file fixtures for LangChain file-input integration tests.

Each generator produces a file containing the same known invoice payload (see
``INVOICE_TEXT``). The generated files are checked into
``tests/langchain/fixtures/files/`` so test runs are deterministic and don't
re-render binaries on every invocation. The generators are kept so we can
regenerate the bundle when adding a new format or tweaking content — run::

    python -m tests.langchain.file_fixtures

Tests load files via ``load_fixture(fmt)`` and convert them into a LangChain
``HumanMessage`` content list via ``build_human_content_block(fmt, data)``,
which assertion code then sends through ``model.with_structured_output``.
"""

from __future__ import annotations

import base64
import csv
import io
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

INVOICE_NUMBER = "INV-7421"
CUSTOMER = "Acme Corp"
# Whole-dollar amount on purpose: when rendered into an image, OCR sometimes
# drops the decimal point of values like "$1234.56" and the model then returns
# 123456.0 instead of 1234.56 (observed on Bedrock-hosted Claude and flaky
# Gemini thinking runs). An integer total sidesteps the ambiguity.
TOTAL_AMOUNT = 4200
DUE_DATE = "2026-03-15"

INVOICE_TEXT = (
    f"Invoice Number: {INVOICE_NUMBER}\n"
    f"Customer: {CUSTOMER}\n"
    f"Total Amount: ${TOTAL_AMOUNT}\n"
    f"Due Date: {DUE_DATE}\n"
)

FIXTURES_DIR = Path(__file__).parent / "fixtures" / "files"


class InvoiceInfo(BaseModel):
    """Structured output schema returned by the model under test."""

    invoice_number: str = Field(description="The invoice number, e.g. INV-1234")
    customer: str = Field(description="The customer name on the invoice")
    total_amount: float = Field(description="The total amount in USD as a number")


def make_txt() -> bytes:
    return INVOICE_TEXT.encode("utf-8")


def make_md() -> bytes:
    body = (
        "# Invoice\n\n"
        f"- **Invoice Number:** {INVOICE_NUMBER}\n"
        f"- **Customer:** {CUSTOMER}\n"
        f"- **Total Amount:** ${TOTAL_AMOUNT}\n"
        f"- **Due Date:** {DUE_DATE}\n"
    )
    return body.encode("utf-8")


def make_csv() -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["field", "value"])
    writer.writerow(["invoice_number", INVOICE_NUMBER])
    writer.writerow(["customer", CUSTOMER])
    writer.writerow(["total_amount", str(TOTAL_AMOUNT)])
    writer.writerow(["due_date", DUE_DATE])
    return buf.getvalue().encode("utf-8")


def make_html() -> bytes:
    body = (
        "<!DOCTYPE html><html><body>"
        "<h1>Invoice</h1>"
        f"<p>Invoice Number: <strong>{INVOICE_NUMBER}</strong></p>"
        f"<p>Customer: {CUSTOMER}</p>"
        f"<p>Total Amount: ${TOTAL_AMOUNT}</p>"
        f"<p>Due Date: {DUE_DATE}</p>"
        "</body></html>"
    )
    return body.encode("utf-8")


def make_pdf() -> bytes:
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.setFont("Helvetica", 14)
    y = 800
    for line in INVOICE_TEXT.splitlines():
        c.drawString(72, y, line)
        y -= 24
    c.save()
    return buf.getvalue()


def make_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Invoice", level=1)
    for line in INVOICE_TEXT.splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Invoice"
    ws.append(["field", "value"])
    ws.append(["invoice_number", INVOICE_NUMBER])
    ws.append(["customer", CUSTOMER])
    ws.append(["total_amount", TOTAL_AMOUNT])
    ws.append(["due_date", DUE_DATE])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_image(fmt: str) -> bytes:
    """Render the invoice text into an image so the model has to OCR it."""
    from PIL import Image, ImageDraw, ImageFont

    width, height = 640, 360
    img = Image.new("RGB", (width, height), color="white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 22)
    except OSError:
        font = ImageFont.load_default()
    y = 40
    for line in INVOICE_TEXT.splitlines():
        draw.text((40, y), line, fill="black", font=font)
        y += 40
    buf = io.BytesIO()
    save_kwargs: dict[str, Any] = {"format": fmt.upper()}
    if fmt.lower() in ("jpg", "jpeg"):
        save_kwargs["format"] = "JPEG"
        save_kwargs["quality"] = 95
    img.save(buf, **save_kwargs)
    return buf.getvalue()


def make_png() -> bytes:
    return _make_image("PNG")


def make_jpg() -> bytes:
    return _make_image("JPEG")


def make_gif() -> bytes:
    return _make_image("GIF")


def make_webp() -> bytes:
    return _make_image("WEBP")


GENERATORS = {
    "txt": (make_txt, "text/plain"),
    "md": (make_md, "text/markdown"),
    "csv": (make_csv, "text/csv"),
    "html": (make_html, "text/html"),
    "pdf": (make_pdf, "application/pdf"),
    "docx": (make_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "xlsx": (make_xlsx, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    "png": (make_png, "image/png"),
    "jpg": (make_jpg, "image/jpeg"),
    "gif": (make_gif, "image/gif"),
    "webp": (make_webp, "image/webp"),
}

TEXT_LIKE_FORMATS = {"txt", "md", "csv", "html"}
IMAGE_FORMATS = {"png", "jpg", "gif", "webp"}
PDF_FORMATS = {"pdf"}
OFFICE_FORMATS = {"docx", "xlsx"}


def fixture_path(fmt: str) -> Path:
    return FIXTURES_DIR / f"invoice.{fmt}"


def load_fixture(fmt: str) -> bytes:
    """Read the committed fixture for ``fmt`` from ``tests/langchain/fixtures/files``."""
    return fixture_path(fmt).read_bytes()


def regenerate_fixtures() -> None:
    """Re-generate every committed fixture file. Run via ``python -m tests.langchain.file_fixtures``."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    for fmt, (generator, _) in GENERATORS.items():
        fixture_path(fmt).write_bytes(generator())


def extract_text_for_office(fmt: str, data: bytes) -> str:
    """Return readable text extracted from a docx/xlsx blob.

    Most LLM providers don't accept raw Office binaries as content blocks, so
    we pre-extract on the client side and send the result as plain text.
    """
    buf = io.BytesIO(data)
    if fmt == "docx":
        from docx import Document

        doc = Document(buf)
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    if fmt == "xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(buf, read_only=True, data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                lines.append(",".join("" if v is None else str(v) for v in row))
        return "\n".join(lines)
    raise ValueError(f"extract_text_for_office does not handle {fmt!r}")


def build_human_content_block(fmt: str, data: bytes) -> list[dict[str, Any]]:
    """Return the LangChain HumanMessage content blocks for a generated file.

    - text-like formats are sent inline as a fenced code block in a text block
    - images go through the standard ``image`` block (base64)
    - pdf goes through the standard ``file`` block (base64)
    - docx/xlsx are pre-extracted to text and sent as a text block
    """
    prompt = "Extract the invoice number, customer, and total amount from the attached file."
    if fmt in TEXT_LIKE_FORMATS:
        text = data.decode("utf-8")
        return [
            {"type": "text", "text": f"{prompt}\n\n```{fmt}\n{text}\n```"},
        ]
    if fmt in OFFICE_FORMATS:
        text = extract_text_for_office(fmt, data)
        return [
            {"type": "text", "text": f"{prompt}\n\n```{fmt}\n{text}\n```"},
        ]
    if fmt in PDF_FORMATS:
        return [
            {"type": "text", "text": prompt},
            {
                "type": "file",
                "base64": base64.b64encode(data).decode("ascii"),
                "mime_type": "application/pdf",
            },
        ]
    if fmt in IMAGE_FORMATS:
        mime = GENERATORS[fmt][1]
        return [
            {"type": "text", "text": prompt},
            {
                "type": "image",
                "base64": base64.b64encode(data).decode("ascii"),
                "mime_type": mime,
            },
        ]
    raise ValueError(f"Unsupported fixture format: {fmt!r}")


if __name__ == "__main__":
    regenerate_fixtures()
    print(f"Wrote {len(GENERATORS)} fixtures to {FIXTURES_DIR}")
